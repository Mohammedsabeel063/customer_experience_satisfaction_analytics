import streamlit as st
import pandas as pd
import numpy as np
import plotly.express as px
import re, string, io, tempfile
import matplotlib.pyplot as plt
from wordcloud import WordCloud
from datetime import datetime

# ----------------- NLP imports -----------------
import nltk
from nltk.corpus import stopwords, wordnet
from nltk.stem import WordNetLemmatizer
from nltk.sentiment import SentimentIntensityAnalyzer
from gensim import corpora, models

# Optional: GenAI
import openai

# Report exports
from fpdf import FPDF
from docx import Document
from docx.shared import Inches
from PIL import Image

# ----------------- One-time NLTK downloads -----------------
@st.cache_resource
def _download_nltk():
    for resource in ["stopwords", "vader_lexicon", "wordnet", "punkt"]:
        try:
            nltk.data.find(f"corpora/{resource}")
        except:
            nltk.download(resource)
_download_nltk()

# ----------------- Page setup -----------------
st.set_page_config(page_title="Customer Experience & Satisfaction Analytics", layout="wide")
st.title("📊 Customer Experience & Satisfaction Analytics")
st.caption("Advanced NLP • Sentiment & Topic Modeling • GenAI Summaries • Hackathon-ready Dashboard")

# ----------------- Load data -----------------
@st.cache_data
def load_data(path: str) -> pd.DataFrame:
    df = pd.read_excel(path)
    colmap = {c.lower(): c for c in df.columns}
    canonical = {
        "feedback": ["customerfeedback", "feedback", "review", "comment", "text"],
        "service": ["service", "service type", "internetservice", "feature"],
        "gender": ["gender"],
        "location": ["location", "region", "city"],
        "age": ["age", "agegroup"],
        "date": ["date", "feedbackdate", "created_at", "timestamp"],
    }
    def pick(name_list, new_name):
        for nm in name_list:
            if nm in colmap:
                df[new_name] = df[colmap[nm]]
                return
    pick(canonical["feedback"], "Feedback")
    pick(canonical["service"], "Service")
    pick(canonical["gender"], "Gender")
    pick(canonical["location"], "Location")
    pick(canonical["age"], "Age")
    pick(canonical["date"], "Date")
    if "Date" in df.columns:
        df["Date"] = pd.to_datetime(df["Date"], errors="coerce")
    return df

DATA_PATH = "customer exp.xlsx"
df = load_data(DATA_PATH)

if "Feedback" not in df.columns:
    st.error("Could not find a feedback text column.")
    st.stop()

# ----------------- Text preprocessing -----------------
lemmatizer = WordNetLemmatizer()
stop_words = set(stopwords.words("english"))
punct_tbl = str.maketrans("", "", string.punctuation)

def normalize_token(tok):
    tok = tok.lower()
    tok = tok.translate(punct_tbl)
    tok = re.sub(r"\d+", "", tok)
    return tok

def lemmatize_word(word):
    return lemmatizer.lemmatize(word, pos=wordnet.NOUN)

def clean_text(text):
    if not isinstance(text, str): 
        return ""
    tokens = nltk.word_tokenize(text)
    tokens = [normalize_token(t) for t in tokens if t.strip()]
    tokens = [t for t in tokens if t and t not in stop_words and len(t) > 2]
    tokens = [lemmatize_word(t) for t in tokens]
    return tokens

@st.cache_data
def build_clean_cols(_df: pd.DataFrame):
    _df = _df.copy()
    _df["CleanTokens"] = _df["Feedback"].astype(str).apply(clean_text)
    _df["CleanText"] = _df["CleanTokens"].apply(lambda toks: " ".join(toks))
    return _df

df = build_clean_cols(df)

# ----------------- Sentiment -----------------
@st.cache_resource
def get_vader():
    return SentimentIntensityAnalyzer()
sia = get_vader()

@st.cache_data
def compute_sentiment(_df: pd.DataFrame):
    _df = _df.copy()
    scores = _df["Feedback"].astype(str).apply(lambda t: sia.polarity_scores(t)["compound"])
    _df["SentimentScore"] = scores
    _df["Sentiment"] = _df["SentimentScore"].apply(lambda s: "Positive" if s>=0.05 else ("Negative" if s<=-0.05 else "Neutral"))
    return _df

df = compute_sentiment(df)

# ----------------- Topic Modeling -----------------
@st.cache_data
def train_lda(tokens_list, num_topics=7, passes=8, min_freq=5):
    dictionary = corpora.Dictionary(tokens_list)
    dictionary.filter_extremes(no_below=min_freq, no_above=0.5)
    corpus = [dictionary.doc2bow(tokens) for tokens in tokens_list]
    if sum(len(doc) for doc in corpus) == 0:
        return None, None, None
    lda = models.LdaModel(corpus=corpus, id2word=dictionary, num_topics=num_topics, passes=passes, random_state=42, alpha="auto", eta="auto")
    return lda, corpus, dictionary

tokens_list = df["CleanTokens"].tolist()
lda_model, lda_corpus, lda_dict = train_lda(tokens_list)

# ----------------- Sidebar Filters -----------------
st.sidebar.header("Filters")
services = df["Service"].dropna().unique().tolist() if "Service" in df.columns else []
service_sel = st.sidebar.multiselect("Service", options=services, default=services)
genders = df["Gender"].dropna().unique().tolist() if "Gender" in df.columns else []
gender_sel = st.sidebar.multiselect("Gender", options=genders, default=genders)
locations = df["Location"].dropna().unique().tolist() if "Location" in df.columns else []
loc_sel = st.sidebar.multiselect("Location", options=locations, default=locations)
if "Date" in df.columns and df["Date"].notna().any():
    min_d, max_d = df["Date"].min(), df["Date"].max()
    start_d, end_d = st.sidebar.date_input("Date range", value=(min_d, max_d))
else:
    start_d, end_d = None, None

# ----------------- Filter data -----------------
df_f = df.copy()
if services and "Service" in df_f.columns: df_f = df_f[df_f["Service"].isin(service_sel)]
if genders and "Gender" in df_f.columns: df_f = df_f[df_f["Gender"].isin(gender_sel)]
if locations and "Location" in df_f.columns: df_f = df_f[df_f["Location"].isin(loc_sel)]
if start_d and end_d and "Date" in df_f.columns:
    df_f = df_f[(df_f["Date"] >= pd.to_datetime(start_d)) & (df_f["Date"] <= pd.to_datetime(end_d))]
st.write(f"*Showing {len(df_f)} records* after filters.")

# ----------------- KPI Cards -----------------
col1, col2, col3, col4 = st.columns(4)
col1.metric("Total Feedback", len(df_f))
col2.metric("Positive %", f"{(df_f['Sentiment']=='Positive').mean()*100:.1f}%")
col3.metric("Negative %", f"{(df_f['Sentiment']=='Negative').mean()*100:.1f}%")
col4.metric("Avg Sentiment Score", f"{df_f['SentimentScore'].mean():.3f}")
st.divider()

# ----------------- Charts -----------------
# Pie chart: Sentiment
c1, c2 = st.columns(2)
with c1:
    st.subheader("Sentiment Distribution")
    sent_counts = df_f["Sentiment"].value_counts().reset_index()
    sent_counts.columns = ["Sentiment", "Count"]
    fig = px.pie(sent_counts, values="Count", names="Sentiment", hole=0.4,
                 color="Sentiment", color_discrete_map={"Positive":"green","Neutral":"gray","Negative":"red"})
    st.plotly_chart(fig, use_container_width=True)
    st.download_button("Download Pie Chart", fig.to_image(format="png"), file_name="sentiment_distribution.png")
with c2:
    if "Service" in df_f.columns:
        st.subheader("Sentiment by Service")
        fig2 = px.histogram(df_f, x="Service", color="Sentiment", barmode="group",
                            color_discrete_map={"Positive":"green","Neutral":"gray","Negative":"red"})
        st.plotly_chart(fig2, use_container_width=True)
        st.download_button("Download Service Chart", fig2.to_image(format="png"), file_name="sentiment_by_service.png")
# ----------------- Trend over time -----------------
if "Date" in df_f.columns and df_f["Date"].notna().any():
    st.subheader("📈 Sentiment Trend Over Time")
    trend = df_f.groupby(pd.Grouper(key="Date", freq="W"))["SentimentScore"].mean().reset_index()
    fig3 = px.line(trend, x="Date", y="SentimentScore")
    st.plotly_chart(fig3, use_container_width=True)

# ----------------- Topic Modeling -----------------
st.subheader("🧩 Topic Modeling (LDA)")
if lda_model is None:
    st.info("Not enough signal to train topics. Try reducing filters or check cleaned text.")
else:
    num_topics = st.slider("Number of topics", 3, 12, 7)
    tokens_filtered = df_f["CleanTokens"].tolist()
    lda_f, corpus_f, dict_f = train_lda(tokens_filtered, num_topics=num_topics, passes=6, min_freq=3)

    if lda_f is None:
        st.info("Filtered data has too little text for topics.")
    else:
        rows = []
        for t in range(num_topics):
            words = lda_f.show_topic(t, topn=8)
            rows.append({"Topic": f"Topic {t}", "Top terms": ", ".join([w for w,_ in words])})
        st.dataframe(pd.DataFrame(rows))

        def top_topic_for_doc(bow):
            dist = lda_f.get_document_topics(bow, minimum_probability=0)
            if not dist: return None
            return max(dist, key=lambda x: x[1])[0]

        corpus_filtered = [dict_f.doc2bow(toks) for toks in tokens_filtered]
        top_topics = [top_topic_for_doc(b) for b in corpus_filtered]
        df_f = df_f.copy()
        df_f["TopTopic"] = top_topics

        topic_counts = df_f["TopTopic"].value_counts().reset_index()
        topic_counts.columns = ["Topic", "Count"]
        figt = px.bar(topic_counts, x="Topic", y="Count", title="Topic Volume")
        st.plotly_chart(figt, use_container_width=True)

        topic_sent = df_f.groupby("TopTopic")["SentimentScore"].mean().reset_index()
        topic_sent["Label"] = topic_sent["SentimentScore"].apply(lambda s: "Positive" if s>0 else ("Negative" if s<0 else "Neutral"))
        st.subheader("📌 Topic Sentiment Overview")
        st.dataframe(topic_sent.sort_values("SentimentScore", ascending=False))

        st.markdown("*Sample feedback per topic*")
        for t in sorted(df_f["TopTopic"].dropna().unique().tolist())[:min(5, num_topics)]:
            st.markdown(f"*Topic {t}* ({topic_sent[topic_sent['TopTopic']==t]['Label'].values[0]})")
            ex = df_f[df_f["TopTopic"]==t]["Feedback"].head(3).tolist()
            for e in ex:
                st.write(f"• {e}")

# ----------------- Word Cloud -----------------
st.subheader("Word Cloud")
all_text = " ".join(df_f["CleanText"].astype(str))
if all_text.strip():
    wc = WordCloud(width=1000, height=400, background_color="white").generate(all_text)
    figw, axw = plt.subplots(figsize=(10,4))
    axw.imshow(wc, interpolation="bilinear")
    axw.axis("off")
    st.pyplot(figw, use_container_width=True)
    buf = io.BytesIO()
    figw.savefig(buf, format="png")
    st.download_button("Download WordCloud", buf, file_name="wordcloud.png", mime="image/png")

# ----------------- Segment Summary (Heuristic) -----------------

# ----------------- GenAI Summary -----------------
st.subheader("GenAI Insights")
if st.button("Generate AI Insights"):
    with st.spinner("Generating..."):
        try:
            openai.api_key = st.secrets["OPENAI_API_KEY"]
            feedback_text = " ".join(df_f["Feedback"].dropna().astype(str))[:3000]
            response = openai.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role":"system","content":"You are a CX expert"},
                          {"role":"user","content":f"Summarize the following feedback into Positives, Complaints, Recommendations:\n{feedback_text}"}],
                max_tokens=300
            )
            summary = response.choices[0].message.content
            st.session_state["genai_summary"] = summary
            st.success(summary)
        except Exception as e:
            st.error(e)

# ----------------- Next Best Actions -----------------
st.subheader("Next Best Actions")
if "genai_summary" in st.session_state:
    if st.button("Suggest Actions"):
        with st.spinner("AI is generating actionable recommendations..."):
            try:
                openai.api_key = st.secrets["OPENAI_API_KEY"]
                response2 = openai.chat.completions.create(
                    model="gpt-4o-mini",
                    messages=[{"role":"system","content":"You are a CX strategy expert"},
                              {"role":"user","content":f"Provide 3-5 actionable next best actions based on this summary:\n{st.session_state['genai_summary']}"}],
                    max_tokens=250
                )
                st.session_state["next_best_actions"] = response2.choices[0].message.content
                st.success(st.session_state["next_best_actions"])
            except Exception as e:
                st.error(e)
else:
    st.info("Generate AI Insights first.")

# ----------------- Multi-format Report -----------------
import io, tempfile
from fpdf import FPDF
from docx import Document
from docx.shared import Inches
from PIL import Image
import pandas as pd
import streamlit as st

st.subheader("📥 Download Short Professional Report")

# ----------------- Convert charts to images -----------------
def generate_chart_images(chart_list):
    images = []
    for fig, title in chart_list:
        buf = io.BytesIO()
        if hasattr(fig, "to_image"):  # Plotly
            buf.write(fig.to_image(format="png"))
        else:  # Matplotlib
            fig.savefig(buf, format="png", bbox_inches='tight', facecolor='w')
        buf.seek(0)
        images.append((buf, title))
    return images

chart_list = []
for name, t in [('fig', "Sentiment Distribution"), ('fig2', "Sentiment by Service"), 
                ('fig3', "Sentiment Trend Over Time"), ('fig_service', "Avg Sentiment per Service"),
                ('figt', "Topic Volume"), ('figw', "Word Cloud")]:
    if name in locals():
        chart_list.append((locals()[name], t))
images = generate_chart_images(chart_list)

# ----------------- TXT Export (short) -----------------
report_lines = ["Customer Feedback Report", f"Records: {len(df_f)}"]
if "genai_summary" in st.session_state:
    report_lines.append("\nGenAI Summary:\n" + st.session_state["genai_summary"])
if "next_best_actions" in st.session_state:
    report_lines.append("\nNext Best Actions:\n" + st.session_state["next_best_actions"])

# Topic Modeling (max 5 feedback per topic)
if 'TopTopic' in df_f.columns:
    report_lines.append("\nTopic Modeling (LDA):")
    for t in sorted(df_f["TopTopic"].dropna().unique()):
        report_lines.append(f"\nTopic {t}:")
        if 'lda_f' in locals():
            words = lda_f.show_topic(t, topn=5)
            report_lines.append("Top terms: " + ", ".join([w for w,_ in words]))
        sentiment_label = topic_sent[topic_sent['TopTopic']==t]['Label'].values[0] if 'topic_sent' in locals() else ""
        report_lines.append(f"Sentiment: {sentiment_label}")
        feedbacks = df_f[df_f["TopTopic"]==t]["Feedback"].tolist()[:5]  # only 5 examples
        for f in feedbacks:
            report_lines.append(f"- {f}")

report_text = "\n".join(report_lines)
st.download_button(
    label="📥 Download TXT",
    data=report_text.encode("utf-8"),
    file_name="customer_report_short.txt",
    mime="text/plain"
)

# ----------------- Excel Export -----------------
excel_buf = io.BytesIO()
with pd.ExcelWriter(excel_buf, engine="openpyxl") as writer:
    df_f.head(50).to_excel(writer, index=False, sheet_name="Feedback (Top 50)")
    
    summary_text = ""
    if "genai_summary" in st.session_state:
        summary_text += "GenAI Summary:\n" + st.session_state["genai_summary"] + "\n\n"
    if "next_best_actions" in st.session_state:
        summary_text += "Next Best Actions:\n" + st.session_state["next_best_actions"]
    if summary_text:
        df_summary = pd.DataFrame({"Summary & Actions": summary_text.split("\n")})
        df_summary.to_excel(writer, index=False, sheet_name="Summary")

excel_buf.seek(0)
st.download_button(
    label="📥 Download Excel",
    data=excel_buf,
    file_name="customer_report_short.xlsx",
    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
)

# ----------------- Word Export (short) -----------------
doc = Document()
doc.add_heading("Customer Feedback Report", 0)

if 'Sentiment' in df_f.columns:
    table = doc.add_table(rows=2, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Total Feedback"
    hdr_cells[1].text = "Positive %"
    hdr_cells[2].text = "Negative %"
    hdr_cells[3].text = "Avg Sentiment"
    val_cells = table.rows[1].cells
    val_cells[0].text = str(len(df_f))
    val_cells[1].text = f"{(df_f['Sentiment']=='Positive').mean()*100:.1f}%"
    val_cells[2].text = f"{(df_f['Sentiment']=='Negative').mean()*100:.1f}%"
    val_cells[3].text = f"{df_f['SentimentScore'].mean():.3f}"

if 'TopTopic' in df_f.columns:
    doc.add_heading("Topic Sentiment Overview (Top 5 Feedbacks)", level=1)
    for t in sorted(df_f["TopTopic"].dropna().unique()):
        feedbacks = df_f[df_f["TopTopic"]==t]["Feedback"].tolist()[:5]
        sentiment_label = topic_sent[topic_sent['TopTopic']==t]['Label'].values[0] if 'topic_sent' in locals() else ""
        doc.add_heading(f"Topic {t} - {sentiment_label}", level=2)
        doc.add_paragraph("\n".join([str(f) for f in feedbacks]))

# Charts (single per page)
doc.add_heading("Visual Insights", level=1)
for img_buf, title in images:
    doc.add_page_break()
    doc.add_heading(title, level=2)
    img = Image.open(img_buf)
    doc.add_picture(img_buf, width=Inches(5))
    doc.add_paragraph(f"Figure: {title}")

doc_io = io.BytesIO()
doc.save(doc_io)
doc_io.seek(0)
st.download_button(
    label="📥 Download Word (.docx)",
    data=doc_io,
    file_name="customer_report_short.docx",
    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)

# ----------------- PDF Export (short) -----------------
pdf = FPDF(orientation='P', unit='mm', format='A4')
pdf.set_auto_page_break(auto=True, margin=15)

pdf.add_page()
pdf.set_font("Arial", 'B', 24)
pdf.cell(0, 30, "Customer Feedback Report", align='C', ln=1)
pdf.set_font("Arial", '', 16)
pdf.ln(10)

if 'Sentiment' in df_f.columns:
    pdf.set_font("Arial", 'B', 14)
    pdf.cell(0, 10, "Key Metrics", ln=1)
    pdf.set_fill_color(200, 230, 255)
    pdf.set_font("Arial", '', 12)
    pdf.cell(60, 10, "Total Feedback", border=1, fill=True)
    pdf.cell(40, 10, "Positive %", border=1, fill=True)
    pdf.cell(40, 10, "Negative %", border=1, fill=True)
    pdf.cell(50, 10, "Avg Sentiment", border=1, fill=True)
    pdf.ln()
    pdf.cell(60, 10, str(len(df_f)), border=1)
    pdf.cell(40, 10, f"{(df_f['Sentiment']=='Positive').mean()*100:.1f}%", border=1)
    pdf.cell(40, 10, f"{(df_f['Sentiment']=='Negative').mean()*100:.1f}%", border=1)
    pdf.cell(50, 10, f"{df_f['SentimentScore'].mean():.3f}", border=1)
    pdf.ln(15)

# Topic Sentiment Overview (max 5 feedback per topic)
if 'TopTopic' in df_f.columns:
    pdf.set_font("Arial", 'B', 14)
    pdf.cell(0, 10, "Topic Sentiment Overview (Top 5 Feedbacks)", ln=1)
    pdf.set_font("Arial", '', 12)
    for t in sorted(df_f["TopTopic"].dropna().unique()):
        feedbacks = df_f[df_f["TopTopic"]==t]["Feedback"].tolist()[:5]
        sentiment_label = topic_sent[topic_sent['TopTopic']==t]['Label'].values[0] if 'topic_sent' in locals() else ""
        pdf.set_font("Arial", 'B', 12)
        pdf.cell(0, 8, f"Topic {t} - {sentiment_label}", ln=1)
        pdf.set_font("Arial", '', 12)
        pdf.multi_cell(0, 8, "\n".join([str(f) for f in feedbacks]))

        pdf.ln(2)

# Charts (smaller, 1 per page)
for img_buf, title in images:
    pdf.add_page()
    pdf.set_font("Arial", 'B', 12)
    pdf.multi_cell(0, 5, title, align='C')
    img = Image.open(img_buf)
    with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
        img.save(tmp.name)
        pdf.image(tmp.name, x=20, y=pdf.get_y()+5, w=170)

pdf_bytes = pdf.output(dest="S").encode("latin1", errors="replace")
st.download_button(
    label="📥 Download PDF",
    data=pdf_bytes,
    file_name="customer_report_short.pdf",
    mime="application/pdf"
)


